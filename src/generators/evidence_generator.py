"""
evidence_generator.py — Generador de documentos Word de evidencia.

Pipeline:
1. Lee Excel (DiseñoEjecución) → obtiene lista de CPs
2. Escanea Insumos/ → analiza .json, .png, .mp4 (OCR + parsing)
3. Relaciona insumos con CPs mediante matching fuzzy + semántico
4. Copia plantilla.docx con versionado y la llena con evidencia
"""

import json
import re
import logging
import shutil
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from thefuzz import fuzz

from src.data.excel_reader import read_hu_excel_with_details, TestCase
from src.engines.local.ocr_engine import extract_text, is_available as ocr_available
from src.engines.local.frame_extractor import extract_frames_fps, is_available as ffmpeg_available

logger = logging.getLogger(__name__)


def _set_cell_shading(cell, color_hex: str):
    """Aplica color de fondo (shading) a una celda de tabla Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


SPANISH_STOPWORDS = {
    'de', 'la', 'el', 'en', 'es', 'un', 'una', 'para', 'que', 'se', 'con',
    'del', 'los', 'las', 'por', 'al', 'no', 'si', 'su', 'le', 'ya', 'pero',
    'sin', 'como', 'todo', 'esta', 'este', 'cada', 'más', 'fue', 'son', 'hay',
    'has', 'ser', 'sus', 'nos', 'les', 'cual', 'uno', 'dos', 'tres', 'sistema',
}

# Singleton del modelo semántico (se carga una vez)
_semantic_model = None


def _get_semantic_model():
    """Carga el modelo de sentence-transformers (lazy, una sola vez)."""
    global _semantic_model
    if _semantic_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            logger.info("Cargando modelo semántico paraphrase-multilingual-MiniLM-L12-v2...")
            _semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
            logger.info("Modelo semántico cargado.")
        except Exception as e:
            logger.warning(f"No se pudo cargar modelo semántico: {e}")
            _semantic_model = False  # Marcar como fallido
    return _semantic_model if _semantic_model is not False else None


@dataclass
class InsumoFile:
    """Representa un archivo multimedia analizado de Insumos/."""
    path: Path
    file_type: str              # "png", "mp4", "json"
    json_description: str = ""
    json_annotations: list = field(default_factory=list)
    ocr_text: str = ""
    timestamp_info: str = ""
    matched_cp: str = ""
    match_score: float = 0.0


def _clean_text(text: str) -> str:
    """Limpia texto para matching: remueve stopwords y normaliza."""
    if not text:
        return ""
    words = re.findall(r'\b\w+\b', text.lower())
    return " ".join(w for w in words if w not in SPANISH_STOPWORDS)


def _parse_timestamp(description: str) -> str:
    """Extrae info de timestamp del JSON description."""
    match = re.search(r'\[Timestamp:\s*([^\]]+)\]', description)
    if match:
        return match.group(1).strip()
    match = re.search(r'\[Segment:\s*([^\]]+)\]', description)
    if match:
        return match.group(1).strip()
    return ""


def _next_version(hu_path: Path, hu_id: str) -> int:
    """Determina la siguiente versión del archivo de evidencia."""
    existing = list(hu_path.glob(f"Evidencia_{hu_id}_V*.docx"))
    if not existing:
        return 1
    versions = []
    for f in existing:
        m = re.search(r'_V(\d+)\.docx$', f.name)
        if m:
            versions.append(int(m.group(1)))
    return max(versions) + 1 if versions else 1


class EvidenceGenerator:
    """Genera documentos Word de evidencia desde Insumos + Excel."""

    def __init__(self, base_dir: Path):
        self.base_dir = base_dir

    # ── Scan Insumos ────────────────────────────────────────────────────────────

    def scan_insumos(self, hu_path: Path) -> List[InsumoFile]:
        """Escanea la carpeta Insumos/ de una HU y analiza cada archivo."""
        insumos_dir = hu_path / "Insumos"
        if not insumos_dir.exists():
            logger.warning(f"No existe carpeta Insumos/ en {hu_path.name}")
            return []

        files = []
        for f in sorted(insumos_dir.iterdir()):
            if f.is_dir() or f.name.startswith('.'):
                continue
            ext = f.suffix.lower()
            if ext not in ('.png', '.mp4', '.json'):
                continue
            files.append(f)

        insumos = []
        json_map = {}

        for f in files:
            ext = f.suffix.lower()

            if ext == '.json':
                insumo = self._analyze_json(f)
                if insumo:
                    insumos.append(insumo)
                    ts = _parse_timestamp(insumo.json_description)
                    if ts:
                        json_map[ts] = insumo

            elif ext == '.png':
                insumo = self._analyze_png(f)
                insumos.append(insumo)

            elif ext == '.mp4':
                insumo = self._analyze_mp4(f)
                insumos.append(insumo)

        self._associate_by_timestamp(insumos, json_map)

        logger.info(
            f"Insumos en {hu_path.name}: "
            f"{sum(1 for i in insumos if i.file_type == 'png')} png, "
            f"{sum(1 for i in insumos if i.file_type == 'mp4')} mp4, "
            f"{sum(1 for i in insumos if i.file_type == 'json')} json"
        )
        return insumos

    def _analyze_json(self, json_path: Path) -> Optional[InsumoFile]:
        """Analiza un archivo JSON de finding."""
        try:
            content = json_path.read_text(encoding='utf-8').strip()
            if not content:
                return None
            data = json.loads(content)
            return InsumoFile(
                path=json_path,
                file_type="json",
                json_description=data.get("description", ""),
                json_annotations=data.get("annotations", []),
                timestamp_info=_parse_timestamp(data.get("description", "")),
            )
        except (json.JSONDecodeError, Exception) as e:
            logger.warning(f"Error leyendo JSON {json_path.name}: {e}")
            return None

    def _analyze_png(self, png_path: Path) -> InsumoFile:
        """Analiza un archivo PNG con OCR."""
        ocr_text = ""
        if ocr_available():
            try:
                ocr_text = extract_text(png_path, lang="spa")
            except Exception as e:
                logger.warning(f"Error OCR en {png_path.name}: {e}")
        return InsumoFile(
            path=png_path,
            file_type="png",
            ocr_text=ocr_text,
        )

    def _analyze_mp4(self, mp4_path: Path) -> InsumoFile:
        """Analiza un video: extrae frames y ejecuta OCR."""
        if not ffmpeg_available():
            logger.warning("ffmpeg no disponible, saltando análisis de video")
            return InsumoFile(path=mp4_path, file_type="mp4")

        import tempfile
        with tempfile.TemporaryDirectory(prefix="evidence_frames_") as tmp:
            tmp_path = Path(tmp)
            frames = extract_frames_fps(mp4_path, tmp_path, fps=1)

            if not frames:
                return InsumoFile(path=mp4_path, file_type="mp4")

            if len(frames) > 10:
                step = len(frames) // 10
                frames = frames[::step][:10]

            ocr_texts = []
            if ocr_available():
                for frame in frames:
                    try:
                        text = extract_text(frame, lang="spa")
                        if text.strip():
                            ocr_texts.append(text)
                    except Exception:
                        pass

            combined_ocr = " ".join(ocr_texts) if ocr_texts else ""
            return InsumoFile(
                path=mp4_path,
                file_type="mp4",
                ocr_text=combined_ocr,
            )

    def _associate_by_timestamp(self, insumos: List[InsumoFile], json_map: Dict[str, InsumoFile]):
        """Asocia descripciones de JSON a imágenes/video por nombre de archivo."""
        for insumo in insumos:
            if insumo.file_type not in ("png", "mp4"):
                continue
            base_name = insumo.path.stem
            for ts, json_insumo in json_map.items():
                json_base = json_insumo.path.stem.replace("finding_", "")
                insumo_base = base_name.replace("finding_", "")
                if json_base == insumo_base:
                    insumo.json_description = json_insumo.json_description
                    insumo.json_annotations = json_insumo.json_annotations
                    insumo.timestamp_info = json_insumo.timestamp_info
                    break

    # ── Matching: Fuzzy + Semántico ─────────────────────────────────────────────

    def _fuzzy_score(self, target: str, source: str) -> float:
        """Score fuzzy entre dos textos."""
        if not source:
            return 0.0
        return fuzz.token_set_ratio(_clean_text(target), _clean_text(source)) / 100

    def _semantic_score(self, target: str, source: str, model) -> float:
        """Score semántico usando sentence-transformers."""
        if not source or not target:
            return 0.0
        try:
            import torch
            from sentence_transformers import util
            emb_target = model.encode(target, convert_to_tensor=True)
            emb_source = model.encode(source, convert_to_tensor=True)
            return float(util.cos_sim(emb_target, emb_source).item())
        except Exception:
            return 0.0

    def _combined_score(
        self,
        cp: TestCase,
        insumo: InsumoFile,
        semantic_model=None,
    ) -> float:
        """
        Score combinado: fuzzy + semántico.
        Pesos: fuzzy 0.4, semántico 0.6 (si disponible), fallback a fuzzy puro.
        """
        target = f"{cp.nombre} {cp.resultado_esperado}"

        # Fuzzy: JSON (0.4) + OCR (0.3) = 0.7 max
        json_f = self._fuzzy_score(target, insumo.json_description)
        ocr_f = self._fuzzy_score(target, insumo.ocr_text)
        fuzzy_total = 0.5 * json_f + 0.5 * ocr_f

        if semantic_model:
            # Semántico: JSON (0.4) + OCR (0.3) = 0.7 max
            json_s = self._semantic_score(target, insumo.json_description, semantic_model)
            ocr_s = self._semantic_score(target, insumo.ocr_text, semantic_model)
            semantic_total = 0.5 * json_s + 0.5 * ocr_s
            # Combinar: fuzzy 40%, semántico 60%
            return 0.4 * fuzzy_total + 0.6 * semantic_total

        # Sin semántico: fuzzy puro normalizado a 0-1
        return fuzzy_total

    def match_insumos_to_cps(
        self,
        insumos: List[InsumoFile],
        test_cases: List[TestCase],
    ) -> Dict[str, List[InsumoFile]]:
        """
        Relaciona insumos con CPs usando matching fuzzy + semántico.

        Returns:
            Dict mapeando cp_id → lista de InsumoFile ordenada por score
        """
        multimedia = [i for i in insumos if i.file_type in ("png", "mp4")]
        if not multimedia or not test_cases:
            return {}

        semantic_model = _get_semantic_model()

        mapping = {tc.id: [] for tc in test_cases}

        for insumo in multimedia:
            best_cp = None
            best_score = 0.0

            for tc in test_cases:
                score = self._combined_score(tc, insumo, semantic_model)
                if score > best_score:
                    best_score = score
                    best_cp = tc.id

            if best_cp:
                insumo.matched_cp = best_cp
                insumo.match_score = best_score
                mapping[best_cp].append(insumo)

        for cp_id in mapping:
            mapping[cp_id].sort(key=lambda x: x.match_score, reverse=True)

        matched_count = sum(1 for cp_id, files in mapping.items() if files)
        method = "fuzzy+semántico" if semantic_model else "fuzzy"
        logger.info(
            f"Matching ({method}): {matched_count}/{len(test_cases)} CPs con al menos 1 insumo"
        )
        return mapping

    # ── Generación de Word (copia plantilla con versionado) ─────────────────────

    def _find_next_version(self, hu_path: Path, hu_id: str) -> int:
        """Busca la siguiente versión disponible."""
        return _next_version(hu_path, hu_id)

    def generate_word(
        self,
        hu_path: Path,
        test_cases: List[TestCase],
        mapping: Dict[str, List[InsumoFile]],
        template_path: Optional[Path] = None,
    ) -> Optional[Path]:
        """
        Copia plantilla.docx, la llena con evidencia y guarda con versionado.

        Args:
            hu_path: Ruta a la carpeta de la HU
            test_cases: Lista de CPs del Excel
            mapping: Dict cp_id → [InsumoFile] del matching
            template_path: Ruta a plantilla.docx (auto-detect si None)

        Returns:
            Path al archivo generado, o None si falló
        """
        if template_path is None:
            template_path = hu_path / "plantilla.docx"

        if not template_path.exists():
            logger.error(f"No se encontró plantilla.docx en {hu_path.name}")
            return None

        hu_id_match = re.search(r'HU-\d+', hu_path.name)
        hu_id = hu_id_match.group(0) if hu_id_match else hu_path.name

        # 1. Determinar versión y copiar plantilla
        version = self._find_next_version(hu_path, hu_id)
        output_name = f"Evidencia_{hu_id}_V{version}.docx"
        output_path = hu_path / output_name

        shutil.copy2(str(template_path), str(output_path))
        logger.info(f"Plantilla copiada: {output_name} (v{version})")

        # 2. Abrir la copia y modificarla
        doc = Document(str(output_path))

        # 3. Encontrar y eliminar las secciones de evidencia existentes de la plantilla
        # La plantilla tiene patrones: "N. EVIDENCIA" + tabla 3-col + tabla 1x1
        self._clear_template_evidence(doc)

        # 4. Insertar secciones de evidencia para cada CP
        cp_count = 0
        for tc in test_cases:
            cp_count += 1
            cp_insumos = mapping.get(tc.id, [])

            # Párrafo de evidencia
            p = doc.add_paragraph(f"{cp_count}. EVIDENCIA — {tc.id}")
            p.style = doc.styles['Normal']
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # Tabla 3-col
            self._add_table_3col(doc, tc)

            # Tabla 1x1 con imagen(es)
            self._add_image_slot(doc, cp_insumos)

        # 5. Guardar
        try:
            doc.save(str(output_path))
            logger.info(f"✓ Word generado: {output_path} ({cp_count} secciones, v{version})")
            return output_path
        except Exception as e:
            logger.error(f"Error guardando Word: {e}")
            return None

    def _clear_template_evidence(self, doc: Document):
        """
        Elimina TODO desde el primer párrafo 'EVIDENCIA' hasta el final del documento.
        Conserva el <w:sectPr> (propiedades de sección) que python-docx necesita.
        """
        body = doc.element.body

        # Encontrar el primer párrafo que contenga 'EVIDENCIA'
        first_evidence_para = None
        for p in doc.paragraphs:
            if re.match(r'^\d+\.\s*EVIDENCIA', p.text.strip()):
                first_evidence_para = p
                break

        if first_evidence_para is None:
            return

        # Recoger todos los elementos desde ese párrafo hasta el final
        start_element = first_evidence_para._element
        elements_to_remove = []
        found_start = False
        for child in list(body):
            if found_start:
                elements_to_remove.append(child)
            elif child is start_element:
                found_start = True
                elements_to_remove.append(child)

        # Eliminar todo EXCEPTO el sectPr
        for elem in elements_to_remove:
            if elem.tag.endswith('}sectPr'):
                continue  # Preservar propiedades de sección
            body.remove(elem)

    def _add_table_3col(self, doc: Document, cp: TestCase):
        """Agrega una tabla 3-columnas con los datos del CP."""
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        headers = ["No. Test Case", "Nombre del Caso de Prueba", "Resultado Esperado"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            _set_cell_shading(cell, "DCE6F1")  # Azul claro original
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        data = [cp.id, cp.nombre, cp.resultado_esperado]
        for i, value in enumerate(data):
            cell = table.rows[1].cells[i]
            cell.text = value or ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

        return table

    def _add_image_slot(self, doc: Document, insumos: List[InsumoFile], max_images: int = 3):
        """Agrega una tabla 1x1 con imagen(es) del matching."""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F0F4F8")  # Gris-azul claro original

        if not insumos:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.text = "(Sin evidencia visual encontrada)"
            return table

        images_inserted = 0
        for insumo in insumos[:max_images]:
            if insumo.file_type == "png" and insumo.path.exists():
                try:
                    p = cell.paragraphs[0] if images_inserted == 0 else cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(insumo.path), width=Inches(5.5))
                    images_inserted += 1
                except Exception as e:
                    logger.warning(f"Error insertando imagen {insumo.path.name}: {e}")
                    if images_inserted == 0:
                        cell.text = f"(Error insertando imagen: {e})"

            elif insumo.file_type == "mp4":
                p = cell.paragraphs[0] if images_inserted == 0 else cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ts = insumo.timestamp_info or "N/A"
                p.text = f"[Video: {insumo.path.name} | Timestamp: {ts}]"
                images_inserted += 1

        if images_inserted == 0:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.text = "(Sin evidencia visual encontrada)"

        return table

    # ── API de alto nivel ──────────────────────────────────────────────────────

    def process_hu(self, hu_path: Path) -> Optional[Path]:
        """
        Pipeline completo para una HU:
        1. Lee Excel → CPs
        2. Escanea Insumos → archivos multimedia
        3. Matching fuzzy + semántico → relación CP ↔ insumos
        4. Copia plantilla → llena con evidencia → guarda versionado
        """
        logger.info(f"{'='*50}")
        logger.info(f"Procesando: {hu_path.name}")

        test_cases = read_hu_excel_with_details(hu_path)
        if not test_cases:
            logger.warning(f"No se encontraron CPs en {hu_path.name}")
            return None

        insumos = self.scan_insumos(hu_path)

        mapping = self.match_insumos_to_cps(insumos, test_cases)

        result = self.generate_word(hu_path, test_cases, mapping)
        if result:
            matched = sum(1 for cp_id, files in mapping.items() if files)
            logger.info(
                f"✓ {hu_path.name}: {len(test_cases)} CPs, "
                f"{matched} con evidencia, archivo generado"
            )
        return result
